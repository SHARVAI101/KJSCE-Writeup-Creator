from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from google import google
import urllib
import urllib.request
import nltk
from nltk.corpus import stopwords
import re
import os

from sumy.parsers.html import HtmlParser
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.luhn import LuhnSummarizer

import time

import requests
from fake_useragent import UserAgent
from bs4 import BeautifulSoup

print('\n\n~~~~~~~~~~~~~~~~~~~~~~~ THE KJSCE WRITE-UP CREATOR ~~~~~~~~~~~~~~~~~~~~~~~')
print('~~~~~~~~~~~~~~~~~~~~~~~ Created By: Sharvai Patil ~~~~~~~~~~~~~~~~~~~~~~~')

inputfile=input('Enter the name of the writeup to be completed--> ')

codelines=input('Enter the line below which I should add your code --> ')

outputlines=input('Enter the line below which I should add your outputs --> ')

postlablines=input('Enter the line below which Post Lab Subjective questions start --> ')

# importing file
document=Document(inputfile); # importing file
document.save("output.docx"); # new file name for output

# get postlab questions one by one
def getPostLabQ():
	global codelines
	global outputlines
	global postlablines

	flag=False # this flag is set to check if the post lab questions' part has been found or not
	fullText = []
	for para in document.paragraphs:
		# fullText.append(para.text)

		if flag==False:
			if postlablines!='nopostlab' and (para.text.strip()=='Post Lab Subjective Questions' or para.text.strip()=='Post Lab Descriptive Questions' or para.text.strip()==postlablines):
				# print("hey")
				# here we have found the post lab subjective questions' part
				flag=True
			elif codelines!='nocode' and (para.text.strip()=='Implementation Details' or para.text.strip()==codelines):
				# here we have found that the user wants to insert code in his writeup
				path = 'code/'
				fileList = os.listdir(path)
				for i in fileList:
					if i!='readme.txt':
						r=para.add_run('\n'+str(os.path.join(i)))
						r.font.size = Pt(18)
						r.bold = True

						# print(os.path.join(i))
						file = open(os.path.join('code/'+ i), 'r', encoding="utf8")
						# print(file.read())
						para.add_run('\n'+str(file.read()))
						para.alignment = WD_ALIGN_PARAGRAPH.LEFT

			elif outputlines!='nooutput' and (para.text.strip()==outputlines):
				path = 'screenshots/'
				fileList = os.listdir(path)
				for i in fileList:
					# file = open(os.path.join('screenshots/'+ i), 'r')
					if i!='readme.txt':
						para.add_run().add_picture(os.path.join('screenshots/'+ i))
					# print(file.read())

		else:
			# means flag is true and we are in the post lab questions' part
			if(para.text.strip()!=''):
				# means if it is not an empty line (in the word doc)				
				querystring=para.text.strip() # this is our question

				# now we will query the string on google
				datastring=searchGoogle(querystring)
				if datastring != False:
					para.add_run('\nAns.: '+str(datastring))
					# print(datastring)
				else:
					print("An answer to the query was not found in the first ten Google links")

	document.save("output.docx")
	print('\n\n~~~ Write Up has been created and saved in the output.docx file ~~~')

def searchGoogle(querystring):
	# to do -> handle exceptions, re-query on google if there is an exception by going to the next link and same for pdf and ppt
	num_page=1
	linkno=0
	# while(True):

	query = "'"+querystring+"'"
	query = urllib.parse.quote_plus(query) # Format into URL encoding
	number_result = 20

	ua = UserAgent()

	google_url = "https://www.google.com/search?q=" + query + "&num=" + str(number_result)
	response = requests.get(google_url, {"User-Agent": ua.random})
	soup = BeautifulSoup(response.text, "html.parser")

	result_div = soup.find_all('div', attrs = {'class': 'g'})

	links = []
	titles = []
	descriptions = []
	for r in result_div:
		# Checks if each element is present, else, raise exception
		try:
			link = r.find('a', href = True)
			title = r.find('h3', attrs={'class': 'r'}).get_text()
			description = r.find('span', attrs={'class': 'st'}).get_text()
			
			# Check to make sure everything is present before appending
			if link != '' and title != '' and description != '': 
				links.append(link['href'])
				titles.append(title)
				descriptions.append(description)
		# Next loop if one element is not present
		except:
			continue

	to_remove = []
	clean_links = []
	for i, l in enumerate(links):
		clean = re.search('\/url\?q\=(.*)\&sa',l)

		# Anything that doesn't fit the above pattern will be removed
		if clean is None:
			to_remove.append(i)
			continue
		clean_links.append(clean.group(1))

	# Remove the corresponding titles & descriptions
	for x in to_remove:
		del titles[x]
		del descriptions[x]

	# print(clean_links)

	for i in range(len(clean_links)):
		# infinite loop to search for the answer for querystring until it is found
		# loop that tries ten times to search for an answer for the querystring
		try:
			print('~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~')
			print("Question --> "+str(querystring))
			# print(querystring)
			
			# searchresult=google.search(querystring,num_page)

			# searchlink=searchresult[linkno].link # this is the first link of the google search results...we will always go to the first link
			searchlink=clean_links[i]
			print("Search Link --> "+str(searchlink))

			if searchlink[-4:]=='.pdf' or searchlink[-4:]=='.ppt':
				# go to next link id the current link is a ppt or pdf
				print("Can't include ppts or pdfs, trying next link on Google")
				linkno+=1
				if linkno>9:
					# if number of links on one page have been exceede, go to the next google link page
					num_page+=1
					linkno=0
			else:
				LANGUAGE = "english"
				SENTENCES_COUNT = 10

				parser = HtmlParser.from_url(searchlink, Tokenizer(LANGUAGE))

				# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
				# Summarisation using Luhn Summarizer
				stopwords1 = set(stopwords.words('english'))

				datastring=''

				# using the LuhnSummarizer
				summarizer = LuhnSummarizer() 
				summarizer.stop_words = stopwords1
				for sentence in summarizer(parser.document, SENTENCES_COUNT):
					# print(sentence)
					datastring+=str(sentence)

				return datastring
		except:
			linkno+=1
			if linkno>9:
				# if number of links on one page have been exceede, go to the next google link page
				num_page+=1
				linkno=0

		time.sleep(1) #sleep for 10 miliseconds so that Google doesn't throw 503 error


	return False

getPostLabQ()