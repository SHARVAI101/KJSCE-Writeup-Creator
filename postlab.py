from docx import Document
from google import google
import urllib.request
import nltk
from nltk.corpus import stopwords
import re
import os

from sumy.parsers.html import HtmlParser
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.luhn import LuhnSummarizer

print('\n\n~~~~~~~~~~~~~~~~~~~~~~~ THE KJSCE WRITE-UP CREATOR ~~~~~~~~~~~~~~~~~~~~~~~')
print('~~~~~~~~~~~~~~ Created By: Sharvai Patil, Anjali Chaudhari ~~~~~~~~~~~~~~')
print('\n~~~~~~~~~~~~~~ Note: This code is only for solving Post Lab questions ~~~~~~~~~~~~~~')

inputfile=input('Enter the name of the writeup to be completed--> ')

postlablines=input('Enter the line below which Post Lab Subjective questions start --> ')

# importing file
document=Document(inputfile); # importing file
document.save("output.docx"); # new file name for output

# get postlab questions one by one
def getPostLabQ():
	global postlablines

	flag=False # this flag is set to check if the post lab questions' part has been found or not
	fullText = []
	for para in document.paragraphs:
		# fullText.append(para.text)

		if flag==False:
			if para.text.strip()=='Post Lab Subjective Questions' or para.text.strip()=='Post Lab Descriptive Questions' or para.text.strip()==postlablines:
				# print("hey")
				# here we have found the post lab subjective questions' part
				flag=True

		else:
			# means flag is true and we are in the post lab questions' part
			if(para.text.strip()!=''):
				# means if it is not an empty line (in the word doc)				
				querystring=para.text.strip() # this is our question

				# now we will query the string on google
				datastring=searchGoogle(querystring)
				para.add_run('\nAns.: '+str(datastring))
				# print(datastring)

	document.save("output.docx")
	print('\n\n~~~ Write Up has been created and saved in the output.docx file ~~~')

def searchGoogle(querystring):
	# to do -> handle exceptions, re-query on google if there is an exception by going to the next link and same for pdf and ppt
	num_page=1
	linkno=0
	while(True):
		# infinite loop to search for the answer for querystring until it is found
		try:
			print('~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~')
			print("QUERY --> "+str(querystring))
			# print(querystring)
			
			searchresult=google.search(querystring,num_page)

			searchlink=searchresult[linkno].link # this is the first link of the google search results...we will always go to the first link
			print("Search Link --> "+str(searchlink))

			if searchlink[-4:]=='.pdf' or searchlink[-4:]=='.ppt':
				# go to next link id the current link is a ppt or pdf
				print("Can't include ppts or pdfs, trying next link on Google")
				linkno+=1
				if linkno>9:
					# if number of links on one page have been exceede, go to the next google link page
					num_page+=1
					linkno=0
			else:
				LANGUAGE = "english"
				SENTENCES_COUNT = 10

				parser = HtmlParser.from_url(searchlink, Tokenizer(LANGUAGE))

				# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
				# Summarisation using Luhn Summarizer
				stopwords1 = set(stopwords.words('english'))

				datastring=''

				# using the LuhnSummarizer
				summarizer = LuhnSummarizer() 
				summarizer.stop_words = stopwords1
				for sentence in summarizer(parser.document, SENTENCES_COUNT):
					# print(sentence)
					datastring+=str(sentence)

				return datastring
		except:
			linkno+=1
			if linkno>9:
				# if number of links on one page have been exceede, go to the next google link page
				num_page+=1
				linkno=0

	
getPostLabQ()